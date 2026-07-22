from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = PROJECT_ROOT / "reports" / "submission_front_matter"


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)


def title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(11, 37, 69)
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(subtitle)
        sr.italic = True
        sr.font.size = Pt(12)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def build_cover_page() -> None:
    doc = Document()
    configure(doc)
    title(doc, "Research Project Report", "Online Grooming Detection using PASYDA Metadata")
    doc.add_paragraph()
    add_kv_table(
        doc,
        [
            ("Module code", "COMP7029"),
            ("Module name", "Artificial Intelligence and Cyber Security"),
            ("Assessment title", "Research Project Report"),
            ("Assessment type", "Portfolio"),
            ("Dataset", "PASYDA synthetic online grooming metadata dataset"),
            ("Submission date", "[insert submission date]"),
        ],
    )
    doc.add_paragraph("Group members").runs[0].bold = True
    add_kv_table(
        doc,
        [
            ("Member 1", "[insert name] - [insert login]"),
            ("Member 2", "[insert name] - [insert login]"),
            ("Member 3", "[insert name] - [insert login]"),
        ],
    )
    doc.add_paragraph(
        "Submitted documents include the cover page, experiment index, portfolio of experiment reports, four-page summary report, and supporting individual contribution reports where required."
    )
    path = OUT_DIR / "Cover_Page_COMP7029_Grooming_Detection.docx"
    doc.save(path)
    print(f"Wrote {path}")


def build_ai_disclosure() -> None:
    doc = Document()
    configure(doc)
    title(doc, "Generative AI Use Disclosure", "COMP7029 Research Project Report")
    doc.add_paragraph(
        "Generative AI tools were used to support project planning, code structuring, report drafting, wording refinement, and critical review of the methodology and interpretation."
    )
    doc.add_paragraph(
        "The team remained responsible for reviewing the generated content, validating the methodology, checking outputs, interpreting the results, and deciding what to include in the final submission."
    )
    doc.add_paragraph("Use of AI included support for:")
    for item in [
        "summarising the assignment requirements and marking criteria",
        "structuring the pair-level modelling workflow",
        "drafting Python scripts for feature engineering, grouped evaluation, tuning analysis, and report generation",
        "drafting and refining report wording",
        "identifying limitations, false-positive/false-negative risks, and synthetic-data shortcut concerns",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_paragraph(
        "No AI-generated output was submitted without human review. The final interpretation acknowledges that the PASYDA dataset is synthetic and that perfect scores must be treated cautiously."
    )
    doc.add_paragraph(
        "This disclosure should be checked against the latest University of Kent Generative AI assessment guidance before final submission."
    )
    path = OUT_DIR / "Generative_AI_Use_Disclosure.docx"
    doc.save(path)
    print(f"Wrote {path}")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_cover_page()
    build_ai_disclosure()


if __name__ == "__main__":
    main()
