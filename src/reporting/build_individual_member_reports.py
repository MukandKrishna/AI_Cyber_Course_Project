from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = PROJECT_ROOT / "reports" / "individual_member_reports"
METRICS_DIR = PROJECT_ROOT / "outputs" / "metrics" / "grouped_evaluation"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(value)
    run.bold = bold
    run.font.size = Pt(9)
    para.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_text(cell, h, True)
        shade(cell, LIGHT_GRAY)
        cell.width = Inches(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_text(cells[i], value)
            cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def callout(doc: Document, title: str, body: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(9.5)
    body_run = p.add_run(" " + body)
    body_run.font.size = Pt(9.2)
    doc.add_paragraph()


def configure(doc: Document, role: str) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 10, 5),
        ("Heading 2", 12.5, BLUE, 8, 4),
        ("Heading 3", 11, DARK_BLUE, 6, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Individual contribution report - {role}").font.size = Pt(8)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(11, 37, 69)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(10)


def metric(metrics: pd.DataFrame, exp: str, col: str) -> str:
    return f"{float(metrics[metrics['experiment'] == exp].iloc[0][col]):.2f}"


def common_context(doc: Document) -> None:
    doc.add_heading("Project Context", level=1)
    doc.add_paragraph(
        "The group project addressed online grooming detection using the PASYDA synthetic metadata dataset. "
        "Direct inspection showed that the correct supervised unit was a central-node/contact-node pair rather than a raw message row."
    )
    table(
        doc,
        ["Project decision", "Justification"],
        [
            ["Pair-level modelling", "The solution files identify the true pair in each scenario, not labels for individual messages."],
            ["Grouped 5-fold evaluation", "Each fold holds out one complete dataset folder, avoiding folder-level leakage."],
            ["Model portfolio", "Baselines, Random Forest, SVM, and MLP provide meaningful comparison across simple, classical, and deep learning approaches."],
            ["Critical analysis", "No-length ablation and false-alert discussion test whether high scores are robust or shortcut-driven."],
        ],
        [2.0, 4.2],
    )


def member_one(metrics: pd.DataFrame) -> Document:
    doc = Document()
    configure(doc, "Data, EDA, and Feature Engineering Lead")
    add_title(
        doc,
        "Individual Contribution Report - Member 1",
        "Data, EDA, and Feature Engineering Lead | Name: [replace with member name] | Login: [replace]",
    )
    common_context(doc)
    doc.add_heading("Assigned Responsibilities", level=1)
    bullets(
        doc,
        [
            "Inspected PASYDA folder structure and ignored irrelevant archive artefacts such as __MACOSX.",
            "Verified that solution rows were contained in vic_data and vic_data rows were contained in full data files.",
            "Identified the central node in each scenario and generated 11 candidate contact pairs per scenario.",
            "Created the modelling-ready pair-level feature table and enhanced contextual feature table.",
            "Produced EDA plots and report-ready evidence for dataset understanding and feature justification.",
        ],
    )
    doc.add_heading("Deliverables Produced", level=1)
    table(
        doc,
        ["Deliverable", "Purpose"],
        [
            ["outputs/features/pair_level_features_base_frozen.csv", "Frozen base modelling table with 110 pair-level samples."],
            ["outputs/features/pair_level_features_enhanced.csv", "Enhanced table with contextual ranks, z-scores, and sequence features."],
            ["outputs/eda/*.png and *.csv", "EDA visuals and audit tables used to justify modelling decisions."],
            ["src/features/build_enhanced_features.py", "Reproducible script for deriving contextual features from raw PASYDA files."],
        ],
        [2.85, 3.35],
    )
    doc.add_heading("Contribution Quality and Academic Value", level=1)
    callout(
        doc,
        "Main technical contribution:",
        "The data work converted many raw PASYDA message rows into the correct 110 pair-level supervised samples. This prevented incorrect message-level labelling and supported a leakage-safe methodology.",
    )
    doc.add_paragraph(
        "The feature engineering directly supports the marking criteria because it demonstrates insight beyond raw metadata. "
        "The work created behavioural variables for volume, reciprocity, duration, time-of-day activity, short-message behaviour, and scenario-relative ranking."
    )
    doc.add_heading("Reflection", level=1)
    doc.add_paragraph(
        "The most important lesson from this role was that dataset structure must determine the modelling unit. "
        "A raw CSV row looked like a natural sample at first, but the labels were actually pair-level. Recognising this avoided a serious methodological flaw."
    )
    doc.add_paragraph(
        "Future improvement would include adding richer privacy-aware metadata, such as account age, platform context, reporting history, and longitudinal interaction features."
    )
    return doc


def member_two(metrics: pd.DataFrame, baseline: pd.DataFrame) -> Document:
    doc = Document()
    configure(doc, "Classical Machine Learning Lead")
    add_title(
        doc,
        "Individual Contribution Report - Member 2",
        "Classical Machine Learning Lead | Name: [replace with member name] | Login: [replace]",
    )
    common_context(doc)
    doc.add_heading("Assigned Responsibilities", level=1)
    bullets(
        doc,
        [
            "Implemented grouped 5-fold evaluation by dataset folder.",
            "Built heuristic baselines to provide a fair comparison before trained models.",
            "Trained and evaluated Random Forest and SVM models using the same folds and metrics.",
            "Ran no-length ablation to test whether model performance depended on message-length artefacts.",
            "Recorded Top-1 Accuracy, MRR, Precision, Recall, F1, ROC-AUC, and PR-AUC.",
        ],
    )
    doc.add_heading("Key Results", level=1)
    table(
        doc,
        ["Experiment", "Top-1", "MRR", "F1", "Interpretation"],
        [
            ["Highest message count", metric(baseline, "baseline_highest_message_count", "top1_accuracy_mean"), metric(baseline, "baseline_highest_message_count", "mean_reciprocal_rank_mean"), metric(baseline, "baseline_highest_message_count", "f1_at_top1_mean"), "Volume alone failed."],
            ["Lowest mean length", metric(baseline, "baseline_lowest_mean_length", "top1_accuracy_mean"), metric(baseline, "baseline_lowest_mean_length", "mean_reciprocal_rank_mean"), metric(baseline, "baseline_lowest_mean_length", "f1_at_top1_mean"), "Very strong but possible artefact."],
            ["RF contextual", metric(metrics, "random_forest_enhanced", "top1_accuracy_mean"), metric(metrics, "random_forest_enhanced", "mean_reciprocal_rank_mean"), metric(metrics, "random_forest_enhanced", "f1_at_top1_mean"), "Perfect full-feature performance."],
            ["RF no length", metric(metrics, "random_forest_no_length", "top1_accuracy_mean"), metric(metrics, "random_forest_no_length", "mean_reciprocal_rank_mean"), metric(metrics, "random_forest_no_length", "f1_at_top1_mean"), "Performance dropped."],
            ["SVM no length", metric(metrics, "svm_rbf_no_length", "top1_accuracy_mean"), metric(metrics, "svm_rbf_no_length", "mean_reciprocal_rank_mean"), metric(metrics, "svm_rbf_no_length", "f1_at_top1_mean"), "Robust without length features."],
        ],
        [1.8, 0.65, 0.65, 0.65, 2.45],
    )
    doc.add_heading("Contribution Quality and Academic Value", level=1)
    callout(
        doc,
        "Main technical contribution:",
        "The evaluation framework avoided leakage by holding out complete dataset folders and treating folds as internal evaluation repeats, not separate experiments.",
    )
    doc.add_paragraph(
        "The classical ML work strengthened the portfolio by showing that simple baselines, Random Forest, and SVM could be compared under the same protocol. "
        "The no-length ablation added critical analysis by testing shortcut dependence rather than reporting perfect scores uncritically."
    )
    doc.add_heading("Reflection", level=1)
    doc.add_paragraph(
        "The main risk in this role was overinterpreting high scores. The baseline and ablation results showed that perfect accuracy can be misleading if a synthetic feature dominates the task."
    )
    doc.add_paragraph(
        "Future work should include systematic hyperparameter search inside training folds only, calibration analysis, and threshold tuning for false-alert control."
    )
    return doc


def member_three(metrics: pd.DataFrame) -> Document:
    doc = Document()
    configure(doc, "Deep Learning and Final Integration Lead")
    add_title(
        doc,
        "Individual Contribution Report - Member 3",
        "Deep Learning and Final Integration Lead | Name: [replace with member name] | Login: [replace]",
    )
    common_context(doc)
    doc.add_heading("Assigned Responsibilities", level=1)
    bullets(
        doc,
        [
            "Selected a small MLP as the deep learning model because the final data representation is tabular and small.",
            "Trained the MLP using the same grouped folds as the classical models.",
            "Compared MLP results against Random Forest, SVM, and heuristic baselines.",
            "Integrated model comparison, ablation findings, false-alert discussion, and limitations into final report materials.",
            "Prepared the 4-page summary report and experiment portfolio structure.",
        ],
    )
    doc.add_heading("Key Results", level=1)
    table(
        doc,
        ["Experiment", "Top-1", "MRR", "F1", "Interpretation"],
        [
            ["MLP base", metric(metrics, "mlp_small_base", "top1_accuracy_mean"), metric(metrics, "mlp_small_base", "mean_reciprocal_rank_mean"), metric(metrics, "mlp_small_base", "f1_at_top1_mean"), "Perfect on base features."],
            ["MLP contextual", metric(metrics, "mlp_small_enhanced", "top1_accuracy_mean"), metric(metrics, "mlp_small_enhanced", "mean_reciprocal_rank_mean"), metric(metrics, "mlp_small_enhanced", "f1_at_top1_mean"), "Perfect on enhanced features."],
            ["MLP no length", metric(metrics, "mlp_small_no_length", "top1_accuracy_mean"), metric(metrics, "mlp_small_no_length", "mean_reciprocal_rank_mean"), metric(metrics, "mlp_small_no_length", "f1_at_top1_mean"), "Small drop after ablation."],
            ["SVM no length", metric(metrics, "svm_rbf_no_length", "top1_accuracy_mean"), metric(metrics, "svm_rbf_no_length", "mean_reciprocal_rank_mean"), metric(metrics, "svm_rbf_no_length", "f1_at_top1_mean"), "Best robustness result."],
        ],
        [1.65, 0.65, 0.65, 0.65, 2.65],
    )
    doc.add_heading("Contribution Quality and Academic Value", level=1)
    callout(
        doc,
        "Main technical contribution:",
        "The integration work connected the deep learning experiment to the wider evidence base, explaining why an MLP was more appropriate than LSTM, GRU, or Transformer models for this small tabular dataset.",
    )
    doc.add_paragraph(
        "The final integration focused on critical interpretation rather than only reporting accuracy. "
        "This included discussion of false positives, false negatives, synthetic-data artefacts, and human-in-the-loop deployment."
    )
    doc.add_heading("Reflection", level=1)
    doc.add_paragraph(
        "The main lesson from the deep learning role was that model complexity must match data structure. "
        "A sequence model would look more advanced, but it would be poorly justified because the labels and feature table are pair-level and tabular."
    )
    doc.add_paragraph(
        "Future work should test the MLP on larger, externally validated datasets and include calibration analysis so scores can support risk ranking rather than automated accusation."
    )
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    model_metrics = pd.read_csv(METRICS_DIR / "model_metrics_summary.csv")
    baseline_metrics = pd.read_csv(METRICS_DIR / "baseline_metrics_summary.csv")

    reports = [
        ("Member_1_Data_EDA_Feature_Engineering_Report.docx", member_one(model_metrics)),
        ("Member_2_Classical_Machine_Learning_Report.docx", member_two(model_metrics, baseline_metrics)),
        ("Member_3_Deep_Learning_Integration_Report.docx", member_three(model_metrics)),
    ]

    for filename, doc in reports:
        path = OUT_DIR / filename
        doc.save(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
