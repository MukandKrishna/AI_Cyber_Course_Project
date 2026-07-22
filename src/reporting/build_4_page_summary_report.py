from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = PROJECT_ROOT / "reports" / "summary_report"
METRICS_DIR = PROJECT_ROOT / "outputs" / "metrics" / "grouped_evaluation"
EDA_DIR = PROJECT_ROOT / "outputs" / "eda"
DOCX_PATH = REPORT_DIR / "Grooming_Detection_4_Page_Summary_Report.docx"
CHART_PATH = REPORT_DIR / "model_ablation_comparison.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.8)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, LIGHT_GRAY)
        cell.width = Inches(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].width = Inches(widths[idx])

    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        para.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    run.font.size = Pt(9.5)
    para.add_run(" " + body).font.size = Pt(9.2)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.4)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 8, 5),
        ("Heading 2", 12, BLUE, 6, 4),
        ("Heading 3", 10.5, DARK_BLUE, 4, 2),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("COMP7029 PASYDA grooming detection summary").font.size = Pt(8)


def create_ablation_chart() -> None:
    comparison = pd.read_csv(METRICS_DIR / "full_vs_no_length_comparison.csv")
    labels = comparison["model"].tolist()
    x = range(len(labels))

    plt.figure(figsize=(6.5, 2.6))
    plt.bar([i - 0.18 for i in x], comparison["full_top1"], width=0.36, label="Full features", color="#2E74B5")
    plt.bar([i + 0.18 for i in x], comparison["no_length_top1"], width=0.36, label="No length features", color="#8AA6C1")
    plt.ylim(0, 1.08)
    plt.ylabel("Top-1 accuracy")
    plt.xticks(list(x), labels)
    plt.title("Full-feature vs no-length ablation")
    plt.legend(loc="lower right", frameon=False)
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    plt.savefig(CHART_PATH, dpi=180)
    plt.close()


def get_metric(metrics: pd.DataFrame, experiment: str, column: str) -> str:
    row = metrics[metrics["experiment"] == experiment].iloc[0]
    return f"{float(row[column]):.2f}"


def get_metric_pm(metrics: pd.DataFrame, experiment: str, mean_column: str, std_column: str) -> str:
    row = metrics[metrics["experiment"] == experiment].iloc[0]
    return f"{float(row[mean_column]):.2f} +/- {float(row[std_column]):.2f}"


def build_report() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    create_ablation_chart()

    model_metrics = pd.read_csv(METRICS_DIR / "model_metrics_summary.csv")
    baseline_metrics = pd.read_csv(METRICS_DIR / "baseline_metrics_summary.csv")
    tuning_summary = pd.read_csv(METRICS_DIR / "tuning_summary.csv")
    error_summary = pd.read_csv(METRICS_DIR / "no_length_error_summary.csv")

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Detecting Online Grooming in PASYDA Metadata")
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Four-page research project summary: pair-level modelling, grouped evaluation, and critical analysis")
    subrun.font.size = Pt(9.5)
    subrun.italic = True

    doc.add_heading("Page 1: Problem, Dataset, and Task Definition", level=1)
    doc.add_paragraph(
        "This project develops an AI/ML workflow for detecting online grooming on the PASYDA synthetic metadata dataset. "
        "The available data contains communication metadata rather than raw message text, with columns for ID, Source, Destination, Date, and Length."
    )
    add_callout(
        doc,
        "Core methodological decision:",
        "The supervised sample is a central-node/contact-node pair, not an individual message row. The solution files identify the true pair in each scenario, not labels for every message.",
    )
    add_table(
        doc,
        ["Item", "Observed project value"],
        [
            ["Dataset folders", "5 folders: Dataset-1 to Dataset-5"],
            ["Scenarios", "10 scenarios total, 2 per dataset folder"],
            ["Candidate pairs", "11 central-node/contact-node pairs per scenario"],
            ["Positive labels", "1 true grooming-related pair per scenario"],
            ["Final modelling rows", "110 pair-level samples: 10 positive, 100 negative"],
            ["Evaluation grouping", "Hold out one complete dataset folder per fold"],
        ],
        [2.2, 4.1],
    )
    doc.add_paragraph(
        "Training directly on raw message rows would be inappropriate because the labels are pair-level labels. "
        "Instead, every raw CSV file is used to construct behavioural pair features, and the models are trained on the resulting pair-level table."
    )
    doc.add_picture(str(EDA_DIR / "01_dataset_overview.png"), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("Page 2: Feature Engineering and Grouped Evaluation", level=1)
    doc.add_paragraph(
        "Feature engineering converted the original metadata into behavioural indicators for each candidate pair. "
        "The frozen base table contains pair-level interaction statistics, while the enhanced table adds contextual and sequence-derived features."
    )
    add_table(
        doc,
        ["Feature group", "Examples", "Purpose"],
        [
            ["Volume and balance", "message_count, sent_ratio, reciprocity_ratio", "Measure communication intensity and directionality"],
            ["Timing", "duration_hours, evening_ratio, late_night_ratio", "Capture when and how long interactions occur"],
            ["Message length", "mean_length, short_message_ratio", "Test whether brevity is predictive"],
            ["Scenario context", "within-scenario ranks and z-scores", "Normalize pair behaviour against other candidates"],
            ["Sequence behaviour", "first_day_intensity, last_day_intensity, direction_switch_frequency", "Represent interaction rhythm"],
        ],
        [1.55, 2.25, 2.45],
    )
    doc.add_paragraph(
        "The evaluation used grouped 5-fold testing by dataset folder. Each fold trained on four complete folders and tested on the remaining folder, giving 88 training rows and 22 test rows per fold."
    )
    add_callout(
        doc,
        "Class imbalance handling:",
        "The pair-level dataset contains 10 positives and 100 negatives. Random Forest and SVM used balanced class weights; the MLP used positive-class oversampling inside training folds only. Test folds were never oversampled.",
    )
    add_table(
        doc,
        ["Metric", "Why it is used"],
        [
            ["Top-1 Accuracy", "Checks whether the true pair is ranked first in each scenario"],
            ["Mean Reciprocal Rank", "Rewards placing the true pair near the top even when not first"],
            ["Precision / Recall / F1 at Top-1", "Measures false alerts and missed detections at the selected top pair"],
            ["ROC-AUC and PR-AUC", "Evaluate ranking quality under class imbalance"],
        ],
        [1.75, 4.55],
    )
    doc.add_picture(str(EDA_DIR / "02_positive_vs_negative_boxplots.png"), width=Inches(5.85))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("Page 3: Model Comparison and Best Results", level=1)
    doc.add_paragraph(
        "Three model families were trained under the same grouped evaluation protocol: Random Forest, RBF-kernel SVM, and a lightweight MLP. "
        "The MLP was chosen because the final representation is tabular and the labelled dataset is too small for sequence models such as LSTM or Transformers."
    )
    add_callout(
        doc,
        "Most important interpretation:",
        "A single hand-crafted rule, selecting the pair with the shortest average message, achieved the same perfect Top-1 score as the trained full-feature models. This does not mean the models failed; it means PASYDA contains a very strong synthetic signal, so all perfect scores must be interpreted in that context.",
    )
    doc.add_paragraph(
        "Hyperparameters were checked using compact nested grouped validation inside each outer training split. Candidate settings for RF depth/leaf size, SVM C values, and MLP hidden sizes/regularisation all retained perfect full-feature test performance, so the final interpretation focuses on robustness rather than chasing higher scores."
    )
    add_table(
        doc,
        ["Experiment", "Top-1 mean +/- std", "MRR", "F1", "ROC-AUC", "PR-AUC"],
        [
            ["RF base", get_metric_pm(model_metrics, "random_forest_base", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "random_forest_base", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "random_forest_base", "f1_at_top1_mean"), get_metric(model_metrics, "random_forest_base", "roc_auc_mean"), get_metric(model_metrics, "random_forest_base", "pr_auc_mean")],
            ["RF contextual", get_metric_pm(model_metrics, "random_forest_enhanced", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "random_forest_enhanced", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "random_forest_enhanced", "f1_at_top1_mean"), get_metric(model_metrics, "random_forest_enhanced", "roc_auc_mean"), get_metric(model_metrics, "random_forest_enhanced", "pr_auc_mean")],
            ["RF no length", get_metric_pm(model_metrics, "random_forest_no_length", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "random_forest_no_length", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "random_forest_no_length", "f1_at_top1_mean"), get_metric(model_metrics, "random_forest_no_length", "roc_auc_mean"), get_metric(model_metrics, "random_forest_no_length", "pr_auc_mean")],
            ["SVM contextual", get_metric_pm(model_metrics, "svm_rbf_enhanced", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "svm_rbf_enhanced", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "svm_rbf_enhanced", "f1_at_top1_mean"), get_metric(model_metrics, "svm_rbf_enhanced", "roc_auc_mean"), get_metric(model_metrics, "svm_rbf_enhanced", "pr_auc_mean")],
            ["SVM no length", get_metric_pm(model_metrics, "svm_rbf_no_length", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "svm_rbf_no_length", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "svm_rbf_no_length", "f1_at_top1_mean"), get_metric(model_metrics, "svm_rbf_no_length", "roc_auc_mean"), get_metric(model_metrics, "svm_rbf_no_length", "pr_auc_mean")],
            ["MLP contextual", get_metric_pm(model_metrics, "mlp_small_enhanced", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "mlp_small_enhanced", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "mlp_small_enhanced", "f1_at_top1_mean"), get_metric(model_metrics, "mlp_small_enhanced", "roc_auc_mean"), get_metric(model_metrics, "mlp_small_enhanced", "pr_auc_mean")],
            ["MLP no length", get_metric_pm(model_metrics, "mlp_small_no_length", "top1_accuracy_mean", "top1_accuracy_std"), get_metric(model_metrics, "mlp_small_no_length", "mean_reciprocal_rank_mean"), get_metric(model_metrics, "mlp_small_no_length", "f1_at_top1_mean"), get_metric(model_metrics, "mlp_small_no_length", "roc_auc_mean"), get_metric(model_metrics, "mlp_small_no_length", "pr_auc_mean")],
        ],
        [1.55, 1.25, 0.72, 0.68, 0.75, 0.75],
    )
    add_callout(
        doc,
        "Best headline result:",
        "RF, SVM, and MLP all achieved perfect performance on base and contextual features. The no-length ablation shows that RF and MLP depend partly on length cues, while SVM remains robust.",
    )
    doc.add_picture(str(CHART_PATH), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Baseline comparison is also important: highest message count scored 0.00 Top-1, longest duration scored 0.30, but lowest mean length and highest short-message ratio scored 1.00."
    )

    doc.add_page_break()
    doc.add_heading("Page 4: Limitations, False Alerts, and Improvements", level=1)
    doc.add_paragraph(
        "The project results are strong on PASYDA, but the report should avoid claiming that grooming is solved in real-world settings. "
        "The dataset is small, synthetic, and contains only 10 positive pair-level samples."
    )
    doc.add_paragraph(
        "Concrete no-length errors were observed: Random Forest missed 3 of 10 true pairs, and MLP missed 1 of 10. The shared difficult case was perp_4 in Dataset-3, where both RF and MLP failed without length-derived cues. This shows why perfect full-feature scores require sceptical interpretation."
    )
    doc.add_paragraph(
        "Further metadata would make the task more realistic and less dependent on synthetic shortcuts. Account age could help identify newly created or throwaway profiles; platform context could distinguish private direct messages from public or group interactions; report status and moderation outcomes could provide stronger weak labels; and longitudinal behavioural annotations could show escalation patterns over time."
    )
    add_table(
        doc,
        ["Risk or limitation", "Impact on interpretation", "Mitigation or future work"],
        [
            ["Synthetic shortcut", "Message-length baselines are perfect, suggesting possible artefacts", "Report no-length ablation and validate on external data"],
            ["False positives", "Wrongly flagged pairs could waste investigator time or harm users", "Use human review and calibrated alert thresholds"],
            ["False negatives", "A true harmful pair could be missed", "Use MRR and recall-focused monitoring, not accuracy alone"],
            ["Limited metadata", "No message content, account age, platform context, or user history", "Collect richer but privacy-aware features"],
            ["Small positive class", "Only 10 positive pair-level samples limits statistical confidence", "Use larger labelled datasets and repeated validation"],
        ],
        [1.45, 2.35, 2.45],
    )
    add_callout(
        doc,
        "Final conclusion:",
        "The proposed pipeline is leakage-safe and performs strongly on PASYDA. However, deployment should be framed as decision support for human review, not automated accusation.",
    )
    add_bullets(
        doc,
        [
            "Use complete-folder grouped evaluation for all future experiments.",
            "Keep pair-level modelling because the labels identify pairs, not individual messages.",
            "Include no-length ablation in the final discussion to show critical analysis.",
            "Request richer metadata such as conversation platform, account age, temporal history, report status, and validated behavioural annotations.",
            "Treat perfect scores as dataset-specific until tested on independent real-world or semi-realistic data.",
        ],
    )

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_report()
